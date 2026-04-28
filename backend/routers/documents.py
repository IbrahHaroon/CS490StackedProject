"""Document Library router (PRD §6 / Sprint 3 stories S3-001 → S3-010).

Architecture:
  Document        — parent record (title, type, status, ownership, timestamps)
  DocumentVersion — immutable content snapshots; current_version_id on Document
                    points at the latest
  DocumentTag     — many tags per document (S3-006 filter)
  JobDocumentLink — N:N between a Job and a specific DocumentVersion (S3-009)

File-storage helpers (PDF/DOCX/text extraction + write-back) live in this module
so the upload + edit workflow keeps its existing semantics.
"""

from __future__ import annotations

import base64
import os
import tempfile
import urllib.parse
from datetime import date as date_class

from docx import Document as DocxDocument
from fastapi import APIRouter, Depends, File, Form, HTTPException, UploadFile, status
from fastapi.responses import FileResponse, Response
from PyPDF2 import PdfReader
from sqlalchemy.orm import Session
from utils import blob_storage

from database import get_db
from database.auth import get_current_user
from database.database import get_settings
from database.models.document import (
    create_document,
    get_document,
    get_documents_for_user,
    hard_delete_document,
    update_document,
)
from database.models.document_tag import (
    add_tag,
    get_tags_for_document,
    remove_tag,
)
from database.models.document_version import (
    create_document_version,
    get_document_version,
    get_versions_for_document,
    restore_version,
)
from database.models.education import get_educations_by_user
from database.models.experience import get_experiences_by_user
from database.models.job import get_job
from database.models.job_document_link import (
    get_links_for_job,
    link_version_to_job,
    unlink,
)
from database.models.profile import get_profile_by_user_id
from database.models.skill import get_skills_for_user
from database.models.user import User
from schemas import (
    DocumentCreate,
    DocumentResponse,
    DocumentTagCreate,
    DocumentTagResponse,
    DocumentUpdate,
    DocumentVersionCreate,
    DocumentVersionResponse,
    JobDocumentLinkCreate,
    JobDocumentLinkResponse,
)

router = APIRouter()

_ROUTERS_DIR = os.path.dirname(os.path.abspath(__file__))
_BACKEND_DIR = os.path.dirname(_ROUTERS_DIR)
# On Vercel the filesystem is read-only; /tmp is the only writable area.
# Uploaded files stored here do NOT persist across function invocations.
_IS_VERCEL = os.environ.get("VERCEL") == "1"
UPLOAD_BASE = "/tmp/uploads" if _IS_VERCEL else os.path.join(_BACKEND_DIR, "uploads")


# --------------------------------------------------------------------------- #
#  File-storage helpers                                                         #
# --------------------------------------------------------------------------- #


def _extract_pdf_content(file_path: str) -> str:
    with open(file_path, "rb") as f:
        pdf_reader = PdfReader(f)
        return "\n".join(p.extract_text() or "" for p in pdf_reader.pages).strip()


def _extract_docx_content(file_path: str) -> str:
    doc = DocxDocument(file_path)
    return "\n".join(p.text for p in doc.paragraphs).strip()


def _write_docx_content(file_path: str, content: str) -> None:
    from docx.shared import Inches, Pt

    doc = DocxDocument()

    # Set margins (1 inch all around)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Add content with proper formatting
    for line in content.split("\n"):
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            run.font.size = Pt(11)
            run.font.name = "Georgia"

    doc.save(file_path)


def _write_resume_docx(file_path: str, content: str) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    doc = DocxDocument()

    # Set margins (1 inch all around)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    lines = content.split("\n")
    i = 0
    first_content = True

    while i < len(lines):
        line = lines[i].strip()
        i += 1

        if not line:
            if not first_content:
                doc.add_paragraph()
            continue

        first_content = False

        # Check if it's a name (first non-empty line)
        if (
            i == 1
            and line
            and not line.startswith(
                ("Phone:", "Email:", "SUMMARY", "EXPERIENCE", "EDUCATION", "SKILLS")
            )
        ):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = "Georgia"
            p.paragraph_format.space_after = Pt(3)
            continue

        # Section headers (all caps lines like SUMMARY, EXPERIENCE, etc.)
        if line.isupper() and len(line.split()) <= 3:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = "Georgia"
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(3)
            continue

        # Job titles or education (pattern: Title at Company or Degree - School)
        if (" at " in line or " - " in line) and not line.startswith(
            ("Location:", "Phone:", "Email:", "GPA:")
        ):
            if " at " in line:
                parts = line.split(" at ", 1)
                p = doc.add_paragraph()
                run = p.add_run(parts[0])
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = "Georgia"
                run = p.add_run(" at ")
                run.font.name = "Georgia"
                run = p.add_run(parts[1])
                run.font.name = "Georgia"
                p.paragraph_format.space_after = Pt(0)
            elif " - " in line:
                parts = line.split(" - ", 1)
                p = doc.add_paragraph()
                run = p.add_run(parts[0])
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = "Georgia"
                run = p.add_run(" - ")
                run.font.name = "Georgia"
                run = p.add_run(parts[1])
                run.italic = True
                run.font.name = "Georgia"
                p.paragraph_format.space_after = Pt(0)
            continue

        # Dates line (e.g., "Jun 2022 - Present")
        if any(
            month in line
            for month in [
                "Jan",
                "Feb",
                "Mar",
                "Apr",
                "May",
                "Jun",
                "Jul",
                "Aug",
                "Sep",
                "Oct",
                "Nov",
                "Dec",
                "Present",
            ]
        ):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.italic = True
            run.font.size = Pt(10)
            run.font.name = "Georgia"
            p.paragraph_format.space_after = Pt(3)
            continue

        # Bullet points (lines starting with dash)
        if line.startswith("- "):
            p = doc.add_paragraph(line[2:], style="List Bullet")
            p_format = p.paragraph_format
            p_format.left_indent = Inches(0.25)
            p_format.space_after = Pt(0)
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.name = "Georgia"
            continue

        # Contact info (Phone, Email, etc.)
        if line.startswith(("Phone:", "Email:", "Location:")):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.size = Pt(10)
            run.font.name = "Georgia"
            p.paragraph_format.space_after = Pt(0)
            continue

        # Default paragraph
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.name = "Georgia"

    doc.save(file_path)


def _write_pdf_content(file_path: str, content: str) -> None:
    from io import BytesIO

    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=72,
        leftMargin=72,
        topMargin=72,
        bottomMargin=72,
    )
    story = []
    styles = getSampleStyleSheet()
    for line in content.split("\n"):
        if line.strip():
            story.append(Paragraph(line, styles["Normal"]))
            story.append(Spacer(1, 0.12 * inch))
    doc.build(story)
    buffer.seek(0)
    with open(file_path, "wb") as f:
        f.write(buffer.read())


def _update_file_content(file_path: str, filename: str, content: str) -> None:
    ext = os.path.splitext(filename)[1].lower()
    if ext in (".txt", ".md"):
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(content)
    elif ext == ".docx":
        _write_docx_content(file_path, content)
    elif ext == ".pdf":
        _write_pdf_content(file_path, content)
    else:
        raise ValueError(f"Unsupported file type for editing: {ext}")


def _read_file(file_path: str, filename: str) -> dict:
    ext = os.path.splitext(filename or "")[1].lower()
    if ext == ".pdf":
        text = _extract_pdf_content(file_path)
        with open(file_path, "rb") as f:
            b64 = base64.b64encode(f.read()).decode("utf-8")
        return {"content": text, "format": "pdf", "binary_data": b64, "editable": True}
    if ext == ".docx":
        return {
            "content": _extract_docx_content(file_path),
            "format": "docx",
            "editable": True,
        }
    if ext in (".txt", ".md"):
        with open(file_path, "r", encoding="utf-8") as f:
            return {"content": f.read(), "format": ext[1:], "editable": True}
    return {
        "content": f"[Unsupported file type: {ext}]",
        "format": "unknown",
        "editable": False,
    }


def _build_upload_path(
    base: str, first_name: str, last_name: str, user_id: int, filename: str
) -> str:
    last_initial = (last_name or "X")[0].upper()
    first_initial = (first_name or "X")[0].upper()
    full_name = f"{first_name or ''} {last_name or ''}".strip() or f"user_{user_id}"
    return os.path.join(
        base, last_initial, first_initial, full_name, str(user_id), filename
    )


def _ensure_owns(document, current_user: User) -> None:
    if document is None:
        raise HTTPException(status_code=404, detail="Document not found")
    if document.user_id != current_user.user_id:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="You do not have permission to access this document",
        )


def _media_type(filename: str) -> str:
    ext = os.path.splitext(filename or "")[1].lower()
    return {
        ".pdf": "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }.get(ext, "application/octet-stream")


def _mime_for_ext(ext: str) -> str:
    return {
        ".pdf": "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".txt": "text/plain",
        ".md": "text/markdown",
    }.get(ext.lower(), "application/octet-stream")


def _filename_from_location(storage_location: str) -> str:
    """Extract filename from either a local path or a blob URL."""
    return urllib.parse.unquote(
        storage_location.rstrip("/").split("/")[-1].split("?")[0]
    )


def _read_bytes(data: bytes, filename: str) -> dict:
    """Parse raw bytes into the same dict shape returned by _read_file."""
    import io

    ext = os.path.splitext(filename or "")[1].lower()
    if ext == ".pdf":
        try:
            text = "\n".join(
                p.extract_text() or "" for p in PdfReader(io.BytesIO(data)).pages
            ).strip()
        except Exception:
            text = ""
        return {
            "content": text,
            "format": "pdf",
            "binary_data": base64.b64encode(data).decode(),
            "editable": True,
        }
    if ext == ".docx":
        doc = DocxDocument(io.BytesIO(data))
        return {
            "content": "\n".join(p.text for p in doc.paragraphs).strip(),
            "format": "docx",
            "editable": True,
        }
    if ext in (".txt", ".md"):
        return {
            "content": data.decode("utf-8", errors="replace"),
            "format": ext[1:],
            "editable": True,
        }
    return {
        "content": f"[Unsupported file type: {ext}]",
        "format": "unknown",
        "editable": False,
    }


def _build_job_context(job) -> str:
    parts = [
        "\nTARGET JOB:",
        f"Title: {job.title}",
        f"Company: {job.company_name}",
    ]
    if job.location:
        parts.append(f"Location: {job.location}")
    if job.salary:
        parts.append(f"Salary: {job.salary}")
    if job.description:
        parts.append(f"Job Description:\n{job.description}")
    if job.years_of_experience is not None:
        parts.append(f"Years of Experience Expected: {job.years_of_experience}")
    if job.company_research_notes:
        parts.append(f"Company Research:\n{job.company_research_notes}")
    return "\n".join(parts)


# --------------------------------------------------------------------------- #
#  Document CRUD (list, read, create, update, archive/restore, hard delete)     #
# --------------------------------------------------------------------------- #


@router.get("/me", response_model=list[DocumentResponse])
def list_my_documents(
    include_archived: bool = False,
    document_type: str | None = None,
    status_filter: str | None = None,
    tag_filter: str | None = None,
    session: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """List the user's documents. Supports filtering by type, status, and tag."""
    docs = get_documents_for_user(
        session, current_user.user_id, include_deleted=include_archived
    )
    if document_type:
        docs = [d for d in docs if d.document_type == document_type]
    if status_filter:
        docs = [d for d in docs if d.status == status_filter]

    result = []
    for d in docs:
        tag_labels = [t.label for t in get_tags_for_document(session, d.document_id)]
        if tag_filter and tag_filter.lower() not in [t.lower() for t in tag_labels]:
            continue
        result.append(
            DocumentResponse(
                document_id=d.document_id,
                user_id=d.user_id,
                title=d.title,
                document_type=d.document_type,
                status=d.status,
                current_version_id=d.current_version_id,
                is_deleted=d.is_deleted,
                created_at=d.created_at,
                updated_at=d.updated_at,
                tags=tag_labels,
            )
        )
    return result


@router.post("", response_model=DocumentResponse, status_code=status.HTTP_201_CREATED)
def create_document_endpoint(
    body: DocumentCreate,
    session: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Create a Document plus an initial v1 DocumentVersion in one shot."""
    doc = create_document(
        session,
        current_user.user_id,
        body.title,
        body.document_type,
        status=body.status or "Draft",
    )
    if body.content or body.storage_location:
        version = create_document_version(
            session,
            doc.document_id,
            storage_location=body.storage_location,
            content=body.content,
            source=body.source or "manual",
        )
        update_document(session, doc.document_id, current_version_id=version.version_id)
    if body.tags:
        for label in body.tags:
            add_tag(session, doc.document_id, label)
    if body.job_id is not None:
        job = get_job(session, body.job_id)
        if job and job.user_id == current_user.user_id:
            current_version_id = doc.current_version_id or (
                create_document_version(
                    session, doc.document_id, source=body.source or "manual"
                ).version_id
            )
            link_version_to_job(
                session,
                job_id=body.job_id,
                version_id=current_version_id,
                role=body.role or doc.document_type,
            )
    return get_document(session, doc.document_id)


@router.get("/{document_id}", response_model=DocumentResponse)
def read_document(
    document_id: int,
    session: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    doc = get_document(session, document_id)
    _ensure_owns(doc, current_user)
    return doc


@router.put("/{document_id}", response_model=DocumentResponse)
def update_document_endpoint(
    document_id: int,
    body: DocumentUpdate,
    session: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Update document metadata only (S3-002, S3-006, S3-008).

    Use POST /{document_id}/versions to write new content.
    """
    doc = get_document(session, document_id)
    _ensure_owns(doc, current_user)
    return update_document(
        session,
        document_id,
        title=body.title,
        document_type=body.document_type,
        status=body.status,
        is_deleted=body.is_deleted,
    )


@router.delete("/{document_id}", status_code=status.HTTP_204_NO_CONTENT)
def hard_delete_document_endpoint(
    document_id: int,
    session: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Hard delete (cascades to versions/tags/links). Prefer PUT with is_deleted=True for archive."""
    doc = get_document(session, document_id)
    _ensure_owns(doc, current_user)
    # Best-effort: remove stored files for all versions
    for version in get_versions_for_document(session, document_id):
        if blob_storage.is_blob_url(version.storage_location):
            blob_storage.delete(version.storage_location)
        elif version.storage_location and os.path.exists(version.storage_location):
            try:
                os.remove(version.storage_location)
            except Exception:
                pass
    hard_delete_document(session, document_id)


@router.post(
    "/{document_id}/duplicate",
    response_model=DocumentResponse,
    status_code=status.HTTP_201_CREATED,
)
def duplicate_document(
    document_id: int,
    session: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """S3-007: duplicate an existing document (copies the current version content and file)."""
    import shutil

    src = get_document(session, document_id)
    _ensure_owns(src, current_user)

    if not src.current_version_id:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Cannot duplicate: source document has no version",
        )

    cv = get_document_version(session, src.current_version_id)
    if not cv:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Cannot duplicate: source version not found",
        )

    # Create the duplicate document
    new_doc = create_document(
        session,
        current_user.user_id,
        f"{src.title} (copy)",
        src.document_type,
        status=src.status,
    )
    session.flush()  # Ensure new_doc.document_id is available

    new_storage_location = None

    if blob_storage.is_blob_url(cv.storage_location):
        # Duplicate the blob under a new name in the same path prefix
        src_filename = _filename_from_location(cv.storage_location)
        name_parts = os.path.splitext(src_filename)
        dup_name = f"{name_parts[0]}_copy{name_parts[1]}"
        parsed = urllib.parse.urlparse(cv.storage_location)
        parent = parsed.path.lstrip("/").rsplit("/", 1)[0]
        dup_pathname = f"{parent}/{dup_name}" if parent else dup_name
        try:
            data = blob_storage.fetch(cv.storage_location)
            new_storage_location = blob_storage.upload(
                dup_pathname, data, _mime_for_ext(name_parts[1])
            )
        except Exception:
            new_storage_location = None
    elif cv.storage_location and os.path.exists(cv.storage_location):
        # Copy the local file to a new path for the duplicate
        src_path = cv.storage_location
        src_dir = os.path.dirname(src_path)
        src_name = os.path.basename(src_path)
        name_parts = os.path.splitext(src_name)
        dup_name = f"{name_parts[0]}_copy{name_parts[1]}"
        new_storage_location = os.path.join(src_dir, dup_name)
        try:
            shutil.copy2(src_path, new_storage_location)
        except Exception:
            new_storage_location = None

    # Create version for the duplicate
    new_version = create_document_version(
        session,
        new_doc.document_id,
        storage_location=new_storage_location,
        content=cv.content,
        source="duplicate",
    )
    session.flush()  # Ensure new_version.version_id is available

    # Link the version to the document
    update_document(
        session, new_doc.document_id, current_version_id=new_version.version_id
    )
    session.commit()

    return get_document(session, new_doc.document_id)


# --------------------------------------------------------------------------- #
#  Versions                                                                     #
# --------------------------------------------------------------------------- #


@router.get("/{document_id}/versions", response_model=list[DocumentVersionResponse])
def list_versions(
    document_id: int,
    session: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    doc = get_document(session, document_id)
    _ensure_owns(doc, current_user)
    return get_versions_for_document(session, document_id)


@router.post(
    "/{document_id}/versions",
    response_model=DocumentVersionResponse,
    status_code=status.HTTP_201_CREATED,
)
def create_new_version(
    document_id: int,
    body: DocumentVersionCreate,
    session: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Append a new version. Sets it as the document's current_version."""
    doc = get_document(session, document_id)
    _ensure_owns(doc, current_user)
    version = create_document_version(
        session,
        document_id,
        storage_location=body.storage_location,
        content=body.content,
        source=body.source or "manual",
    )
    update_document(session, document_id, current_version_id=version.version_id)
    return version


@router.get("/{document_id}/versions/{version_id}/content")
def read_version_content(
    document_id: int,
    version_id: int,
    session: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Read the content of a specific version (not necessarily the current one).

    Mirrors `/documents/{id}/content` but operates on the version named in the
    URL. Returns 404 if the version doesn't exist or doesn't belong to this
    document, 403 if the document isn't owned by the current user.
    """
    doc = get_document(session, document_id)
    _ensure_owns(doc, current_user)
    version = get_document_version(session, version_id)
    if version is None or version.document_id != document_id:
        raise HTTPException(
            status_code=404, detail="Version not found for this document"
        )
    if version.content:
        return {"content": version.content, "format": "text", "editable": True}
    if blob_storage.is_blob_url(version.storage_location):
        try:
            data = blob_storage.fetch(version.storage_location)
            return _read_bytes(data, _filename_from_location(version.storage_location))
        except Exception as e:  # noqa: BLE001
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Failed to fetch file from storage: {e}",
            )
    if version.storage_location and os.path.exists(version.storage_location):
        try:
            return _read_file(
                version.storage_location,
                os.path.basename(version.storage_location),
            )
        except Exception as e:  # noqa: BLE001
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Failed to read file: {e}",
            )
    raise HTTPException(status_code=404, detail="Version has no readable content")


@router.post(
    "/{document_id}/versions/{version_id}/restore",
    response_model=DocumentResponse,
)
def restore_to_version(
    document_id: int,
    version_id: int,
    session: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Set `current_version_id` to an existing prior version of this document.

    Returns the updated Document. 404 if the version doesn't belong to the
    named document, 403 if the document isn't owned by the current user.
    """
    doc = get_document(session, document_id)
    _ensure_owns(doc, current_user)
    updated = restore_version(session, document_id, version_id)
    if updated is None:
        raise HTTPException(
            status_code=404, detail="Version not found for this document"
        )
    return updated


# --------------------------------------------------------------------------- #
#  Content read / edit (operates on the current version)                        #
# --------------------------------------------------------------------------- #


@router.get("/{document_id}/content")
def read_current_content(
    document_id: int,
    session: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    doc = get_document(session, document_id)
    _ensure_owns(doc, current_user)
    if doc.current_version_id is None:
        raise HTTPException(status_code=404, detail="Document has no version yet")
    version = get_document_version(session, doc.current_version_id)
    if version is None:
        raise HTTPException(status_code=404, detail="Current version missing")
    if version.content:
        return {"content": version.content, "format": "text", "editable": True}
    if blob_storage.is_blob_url(version.storage_location):
        try:
            data = blob_storage.fetch(version.storage_location)
            return _read_bytes(data, _filename_from_location(version.storage_location))
        except Exception as e:  # noqa: BLE001
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Failed to fetch file from storage: {e}",
            )
    if version.storage_location and os.path.exists(version.storage_location):
        try:
            return _read_file(
                version.storage_location,
                os.path.basename(version.storage_location),
            )
        except Exception as e:  # noqa: BLE001
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Failed to read file: {e}",
            )
    raise HTTPException(status_code=404, detail="No content available")


@router.put("/{document_id}/content", response_model=DocumentVersionResponse)
def edit_current_content(
    document_id: int,
    body: dict,
    session: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Edit current version content. If the version is file-backed, the file is
    rewritten on disk; otherwise a new version row is appended carrying the
    text content.
    """
    doc = get_document(session, document_id)
    _ensure_owns(doc, current_user)
    content = body.get("content")
    if content is None:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Content field is required",
        )
    current = (
        get_document_version(session, doc.current_version_id)
        if doc.current_version_id
        else None
    )
    if (
        current
        and current.storage_location
        and os.path.exists(current.storage_location)
    ):
        try:
            _update_file_content(
                current.storage_location,
                os.path.basename(current.storage_location),
                content,
            )
        except ValueError as e:
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=str(e))
        # Append a new version row with the in-memory text snapshot for history.
        new_version = create_document_version(
            session,
            document_id,
            storage_location=current.storage_location,
            content=content,
            source="edit",
        )
    else:
        new_version = create_document_version(
            session, document_id, content=content, source="edit"
        )
    update_document(session, document_id, current_version_id=new_version.version_id)
    return new_version


# --------------------------------------------------------------------------- #
#  Download / Export                                                            #
# --------------------------------------------------------------------------- #


@router.get("/{document_id}/download")
def download_document(
    document_id: int,
    session: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Download the current version of a document as a file.

    If storage_location exists, stream the file. Otherwise, generate a DOCX
    on-the-fly from the content text.
    """
    doc = get_document(session, document_id)
    _ensure_owns(doc, current_user)
    if doc.current_version_id is None:
        raise HTTPException(status_code=404, detail="Document has no version yet")
    version = get_document_version(session, doc.current_version_id)
    if version is None:
        raise HTTPException(status_code=404, detail="Current version missing")

    # If stored in Vercel Blob, fetch and stream directly
    if blob_storage.is_blob_url(version.storage_location):
        filename = _filename_from_location(version.storage_location)
        try:
            data = blob_storage.fetch(version.storage_location)
        except Exception as e:  # noqa: BLE001
            raise HTTPException(
                status_code=status.HTTP_502_BAD_GATEWAY,
                detail=f"Failed to fetch file from storage: {e}",
            )
        return Response(
            content=data,
            media_type=_media_type(filename),
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )

    # If file-backed on disk, stream it
    if version.storage_location and os.path.exists(version.storage_location):
        filename = os.path.basename(version.storage_location)
        media = _media_type(filename)
        return FileResponse(
            path=version.storage_location,
            filename=filename,
            media_type=media,
        )

    # Otherwise, generate DOCX from content
    if version.content:
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp_path = tmp.name
        try:
            _write_docx_content(tmp_path, version.content)
            filename = (
                f"{doc.title}.docx" if not doc.title.endswith(".docx") else doc.title
            )
            return FileResponse(
                path=tmp_path,
                filename=filename,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        except Exception as e:  # noqa: BLE001
            if os.path.exists(tmp_path):
                try:
                    os.remove(tmp_path)
                except Exception:
                    pass
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Failed to generate document: {e}",
            )

    raise HTTPException(status_code=404, detail="No content available")


@router.get("/{document_id}/versions/{version_id}/download")
def download_document_version(
    document_id: int,
    version_id: int,
    session: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Download a specific version of a document as a file.

    Same logic as current-version download, but for a specified version.
    """
    doc = get_document(session, document_id)
    _ensure_owns(doc, current_user)
    version = get_document_version(session, version_id)
    if version is None or version.document_id != document_id:
        raise HTTPException(
            status_code=404, detail="Version not found for this document"
        )

    # If stored in Vercel Blob, fetch and stream directly
    if blob_storage.is_blob_url(version.storage_location):
        filename = _filename_from_location(version.storage_location)
        try:
            data = blob_storage.fetch(version.storage_location)
        except Exception as e:  # noqa: BLE001
            raise HTTPException(
                status_code=status.HTTP_502_BAD_GATEWAY,
                detail=f"Failed to fetch file from storage: {e}",
            )
        return Response(
            content=data,
            media_type=_media_type(filename),
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )

    # If file-backed on disk, stream it
    if version.storage_location and os.path.exists(version.storage_location):
        filename = os.path.basename(version.storage_location)
        media = _media_type(filename)
        return FileResponse(
            path=version.storage_location,
            filename=filename,
            media_type=media,
        )

    # Otherwise, generate DOCX from content
    if version.content:
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp_path = tmp.name
        try:
            _write_docx_content(tmp_path, version.content)
            filename = f"{doc.title} (v{version.version_number}).docx"
            if doc.title.endswith(".docx"):
                filename = f"{doc.title[:-5]} (v{version.version_number}).docx"
            return FileResponse(
                path=tmp_path,
                filename=filename,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        except Exception as e:  # noqa: BLE001
            if os.path.exists(tmp_path):
                try:
                    os.remove(tmp_path)
                except Exception:
                    pass
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Failed to generate document: {e}",
            )

    raise HTTPException(status_code=404, detail="No content available")


# --------------------------------------------------------------------------- #
#  Upload (S3-004)                                                              #
# --------------------------------------------------------------------------- #


@router.post(
    "/upload", response_model=DocumentResponse, status_code=status.HTTP_201_CREATED
)
async def upload_document(
    file: UploadFile = File(...),
    document_type: str = Form(...),
    title: str | None = Form(None),
    status_value: str = Form("Draft"),
    session: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    profile = get_profile_by_user_id(session, current_user.user_id)
    if not profile:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="User profile not found — create a profile before uploading documents",
        )
    contents = await file.read()
    ext = os.path.splitext(file.filename)[1].lower()

    if blob_storage.enabled():
        pathname = blob_storage.build_pathname(
            profile.first_name, profile.last_name, current_user.user_id, file.filename
        )
        try:
            dest_path = blob_storage.upload(pathname, contents, _mime_for_ext(ext))
        except Exception as exc:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Failed to upload file to cloud storage: {exc}",
            )
    else:
        dest_path = _build_upload_path(
            UPLOAD_BASE,
            profile.first_name,
            profile.last_name,
            current_user.user_id,
            file.filename,
        )
        os.makedirs(os.path.dirname(dest_path), exist_ok=True)
        with open(dest_path, "wb") as f:
            f.write(contents)

    # Cache extracted text in DB so the document is readable without the file.
    extracted_content = None
    if ext in (".txt", ".md", ".docx"):
        try:
            extracted_content = _read_bytes(contents, file.filename).get("content")
        except Exception:
            pass  # content stays None; storage_location is the fallback

    doc = create_document(
        session,
        current_user.user_id,
        title or file.filename,
        document_type,
        status=status_value,
    )
    version = create_document_version(
        session,
        doc.document_id,
        storage_location=dest_path,
        content=extracted_content,
        source="upload",
    )
    update_document(session, doc.document_id, current_version_id=version.version_id)
    return get_document(session, doc.document_id)


# --------------------------------------------------------------------------- #
#  Tags (S3-006)                                                                #
# --------------------------------------------------------------------------- #


@router.get("/{document_id}/tags", response_model=list[DocumentTagResponse])
def list_tags(
    document_id: int,
    session: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    doc = get_document(session, document_id)
    _ensure_owns(doc, current_user)
    return get_tags_for_document(session, document_id)


@router.post(
    "/{document_id}/tags",
    response_model=DocumentTagResponse,
    status_code=status.HTTP_201_CREATED,
)
def create_tag(
    document_id: int,
    body: DocumentTagCreate,
    session: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    doc = get_document(session, document_id)
    _ensure_owns(doc, current_user)
    tag = add_tag(session, document_id, body.label)
    if tag is None:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST, detail="Empty tag label"
        )
    return tag


@router.delete("/{document_id}/tags/{label}", status_code=status.HTTP_204_NO_CONTENT)
def delete_tag(
    document_id: int,
    label: str,
    session: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    doc = get_document(session, document_id)
    _ensure_owns(doc, current_user)
    if not remove_tag(session, document_id, label):
        raise HTTPException(status_code=404, detail="Tag not found")


# --------------------------------------------------------------------------- #
#  Job ↔ Document linking (S3-009 / S3-010)                                     #
# --------------------------------------------------------------------------- #


@router.post(
    "/links",
    response_model=JobDocumentLinkResponse,
    status_code=status.HTTP_201_CREATED,
)
def create_link(
    body: JobDocumentLinkCreate,
    session: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    job = get_job(session, body.job_id)
    if not job or job.user_id != current_user.user_id:
        raise HTTPException(status_code=404, detail="Job not found")
    version = get_document_version(session, body.version_id)
    if not version:
        raise HTTPException(status_code=404, detail="Version not found")
    parent_doc = get_document(session, version.document_id)
    _ensure_owns(parent_doc, current_user)
    return link_version_to_job(
        session, job_id=body.job_id, version_id=body.version_id, role=body.role
    )


@router.delete("/links/{link_id}", status_code=status.HTTP_204_NO_CONTENT)
def delete_link(
    link_id: int,
    session: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    # Ownership inferred via the underlying job.
    from database.models.job_document_link import JobDocumentLink

    link = session.get(JobDocumentLink, link_id)
    if link is None:
        raise HTTPException(status_code=404, detail="Link not found")
    job = get_job(session, link.job_id)
    if not job or job.user_id != current_user.user_id:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN, detail="Access denied"
        )
    unlink(session, link_id)


@router.get(
    "/links/by-job/{job_id}",
    response_model=list[JobDocumentLinkResponse],
)
def list_links_for_job(
    job_id: int,
    session: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    job = get_job(session, job_id)
    if not job or job.user_id != current_user.user_id:
        raise HTTPException(status_code=404, detail="Job not found")
    return get_links_for_job(session, job_id)


@router.get("/links/by-job/{job_id}/detailed")
def list_links_for_job_detailed(
    job_id: int,
    session: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Enriched version of /links/by-job that includes document title and type."""
    from sqlalchemy import select as sa_select

    from database.models.document import Document
    from database.models.document_version import DocumentVersion
    from database.models.job_document_link import JobDocumentLink

    job = get_job(session, job_id)
    if not job or job.user_id != current_user.user_id:
        raise HTTPException(status_code=404, detail="Job not found")
    rows = session.execute(
        sa_select(
            JobDocumentLink.link_id,
            JobDocumentLink.job_id,
            JobDocumentLink.version_id,
            JobDocumentLink.role,
            JobDocumentLink.linked_at,
            Document.document_id,
            Document.title.label("document_title"),
            Document.document_type,
        )
        .join(DocumentVersion, JobDocumentLink.version_id == DocumentVersion.version_id)
        .join(Document, DocumentVersion.document_id == Document.document_id)
        .where(JobDocumentLink.job_id == job_id)
        .where(Document.is_deleted.is_(False))
    ).all()
    return [
        {
            "link_id": r.link_id,
            "job_id": r.job_id,
            "version_id": r.version_id,
            "role": r.role,
            "linked_at": r.linked_at.isoformat() if r.linked_at else None,
            "document_id": r.document_id,
            "document_title": r.document_title,
            "document_type": r.document_type,
        }
        for r in rows
    ]


@router.get("/links/me")
def list_my_links(
    session: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """All document-job links for the current user, enriched with doc + job metadata."""
    from sqlalchemy import select as sa_select

    from database.models.document import Document
    from database.models.document_version import DocumentVersion
    from database.models.job import Job
    from database.models.job_document_link import JobDocumentLink

    rows = session.execute(
        sa_select(
            JobDocumentLink.link_id,
            JobDocumentLink.job_id,
            JobDocumentLink.version_id,
            JobDocumentLink.role,
            JobDocumentLink.linked_at,
            Document.document_id,
            Document.title.label("document_title"),
            Document.document_type,
            Job.title.label("job_title"),
            Job.company_name,
        )
        .join(DocumentVersion, JobDocumentLink.version_id == DocumentVersion.version_id)
        .join(Document, DocumentVersion.document_id == Document.document_id)
        .join(Job, JobDocumentLink.job_id == Job.job_id)
        .where(Document.user_id == current_user.user_id)
        .where(Document.is_deleted.is_(False))
    ).all()
    return [
        {
            "link_id": r.link_id,
            "job_id": r.job_id,
            "version_id": r.version_id,
            "role": r.role,
            "linked_at": r.linked_at.isoformat() if r.linked_at else None,
            "document_id": r.document_id,
            "document_title": r.document_title,
            "document_type": r.document_type,
            "job_title": r.job_title,
            "company_name": r.company_name,
        }
        for r in rows
    ]


# --------------------------------------------------------------------------- #
#  AI generation (creates a Document + version + optional job link)             #
# --------------------------------------------------------------------------- #


def _build_profile_text(session: Session, user_id: int) -> str:
    profile = get_profile_by_user_id(session, user_id)
    if not profile:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="User profile not found — create a profile first",
        )
    experiences = get_experiences_by_user(session, user_id)
    educations = get_educations_by_user(session, user_id)
    skills = get_skills_for_user(session, user_id)

    lines = [f"Name: {profile.first_name} {profile.last_name}"]
    if profile.phone_number:
        lines.append(f"Phone: {profile.phone_number}")
    if profile.summary:
        lines.append(f"Summary: {profile.summary}")
    if experiences:
        lines.append("\nWORK EXPERIENCE:")
        for exp in experiences:
            end = exp.end_date.strftime("%b %Y") if exp.end_date else "Present"
            lines.append(
                f"  {exp.title} at {exp.company} ({exp.start_date.strftime('%b %Y')} - {end})"
            )
            if exp.location:
                lines.append(f"  Location: {exp.location}")
            if exp.description:
                lines.append(f"  {exp.description}")
    if educations:
        lines.append("\nEDUCATION:")
        for edu in educations:
            field = f" in {edu.field_of_study}" if edu.field_of_study else ""
            lines.append(f"  {edu.degree}{field} — {edu.school}")
            if edu.gpa:
                lines.append(f"  GPA: {edu.gpa}")
            if edu.start_date and edu.end_date:
                lines.append(f"  {edu.start_date.year} - {edu.end_date.year}")
    if skills:
        lines.append("\nSKILLS:")
        lines.append("  " + ", ".join(s.name for s in skills))
    return "\n".join(lines)


def _call_openai(system_prompt: str, user_message: str) -> str:
    import openai

    settings = get_settings()
    api_key = os.environ.get("OPENAI_API_KEY") or settings.openai_api_key
    if not api_key:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="OpenAI API key is not configured.",
        )
    try:
        client = openai.OpenAI(api_key=api_key)
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_message},
            ],
            max_tokens=4096,
            temperature=0.7,
        )
        return response.choices[0].message.content.strip()
    except openai.AuthenticationError:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Invalid OpenAI API key.",
        )
    except openai.RateLimitError:
        raise HTTPException(
            status_code=status.HTTP_429_TOO_MANY_REQUESTS,
            detail="OpenAI rate limit reached. Try again shortly.",
        )
    except Exception as e:  # noqa: BLE001
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"AI generation failed: {e}",
        )


def _generate_doc(
    session: Session,
    *,
    user: User,
    document_type: str,
    role: str,
    title_prefix: str,
    system_prompt: str,
    body: dict,
) -> dict:
    profile_text = _build_profile_text(session, user.user_id)
    job_id = body.get("job_id")
    job_context = ""
    job_label = ""
    job = None
    if job_id:
        job = get_job(session, job_id)
        if not job or job.user_id != user.user_id:
            raise HTTPException(status_code=404, detail="Job not found")
        job_context = _build_job_context(job)
        job_label = f" - {job.title} @ {job.company_name}"

    instructions = (body.get("instructions") or "").strip()
    user_message = f"Generate from this profile:\n\n{profile_text}"
    if job_context:
        user_message += f"\n\n{job_context}"
    if instructions:
        user_message += f"\n\nAdditional instructions: {instructions}"

    content = _call_openai(system_prompt, user_message)

    doc_title = f"{title_prefix}{job_label} - {date_class.today().isoformat()}"
    doc = create_document(
        session, user.user_id, doc_title, document_type, status="Draft"
    )

    profile = get_profile_by_user_id(session, user.user_id)
    filename = f"{doc_title.replace('/', '_').replace(':', '-')}.docx"

    if blob_storage.enabled():
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            write_path = tmp.name
        if document_type == "Resume":
            _write_resume_docx(write_path, content)
        else:
            _write_docx_content(write_path, content)
        blob_pathname = blob_storage.build_pathname(
            profile.first_name if profile else "",
            profile.last_name if profile else "",
            user.user_id,
            filename,
        )
        with open(write_path, "rb") as f:
            docx_bytes = f.read()
        try:
            os.remove(write_path)
        except Exception:
            pass
        storage_loc = blob_storage.upload(
            blob_pathname,
            docx_bytes,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    else:
        storage_loc = _build_upload_path(
            UPLOAD_BASE,
            profile.first_name if profile else "",
            profile.last_name if profile else "",
            user.user_id,
            filename,
        )
        os.makedirs(os.path.dirname(storage_loc), exist_ok=True)
        if document_type == "Resume":
            _write_resume_docx(storage_loc, content)
        else:
            _write_docx_content(storage_loc, content)

    version = create_document_version(
        session,
        doc.document_id,
        storage_location=storage_loc,
        content=content,
        source="ai",
    )
    update_document(session, doc.document_id, current_version_id=version.version_id)
    if job is not None:
        link_version_to_job(
            session, job_id=job.job_id, version_id=version.version_id, role=role
        )
    return {
        "document_id": doc.document_id,
        "version_id": version.version_id,
        "content": content,
        "title": doc_title,
    }


@router.post("/generate-resume")
def generate_resume(
    body: dict,
    session: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    return _generate_doc(
        session,
        user=current_user,
        document_type="Resume",
        role="resume",
        title_prefix="AI Resume",
        system_prompt=(
            "You are an expert resume writer. Produce a clean ATS-friendly resume from the "
            "candidate profile and (when given) the target job. Plain text. Standard sections "
            "(CONTACT, SUMMARY, EXPERIENCE, EDUCATION, SKILLS). No code fences or preamble."
        ),
        body=body,
    )


@router.post("/generate-cover-letter")
def generate_cover_letter(
    body: dict,
    session: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    return _generate_doc(
        session,
        user=current_user,
        document_type="Cover Letter",
        role="cover_letter",
        title_prefix="AI Cover Letter",
        system_prompt=(
            "You are an expert cover-letter writer. Produce a personalized 3-4 paragraph "
            "cover letter for the target job using the candidate profile. Address it to the "
            "hiring team at the company. End with a clear call to action and the candidate's "
            "name. Plain text only — no preamble or markdown."
        ),
        body=body,
    )


# --------------------------------------------------------------------------- #
#  Backwards-compatible aliases (DocumentCreate shape used by older callers)    #
# --------------------------------------------------------------------------- #


@router.post("/", response_model=DocumentResponse, status_code=status.HTTP_201_CREATED)
def create_document_legacy_path(
    body: DocumentCreate,
    session: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Trailing-slash alias of POST /. Retained for the existing frontend client."""
    return create_document_endpoint(body, session, current_user)
